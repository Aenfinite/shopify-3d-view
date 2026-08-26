import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = docx.Document()
    
    # Page setup - Normal Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.different_first_page_header_footer = True
        
        # Header / Footer for subsequent pages
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Aenfinite Custom Tailoring | Material Management & Shopify Workflow")
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(128, 128, 128)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Confidential — For Internal Technical & Production Use Only")
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(140, 140, 140)

    # Color Palette Constants
    PRIMARY_HEX = "1E3A8A"      # Navy / Deep Blue
    PRIMARY_RGB = RGBColor(30, 58, 138)
    SECONDARY_HEX = "0D9488"    # Teal Accent
    SECONDARY_RGB = RGBColor(13, 148, 136)
    DARK_TEXT_RGB = RGBColor(31, 41, 55)
    MUTED_TEXT_RGB = RGBColor(75, 85, 99)
    BG_LIGHT_HEX = "F8FAFC"     # Slate 50
    BORDER_HEX = "CBD5E1"       # Slate 300

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    def add_callout(text, title="NOTE / IMPORTANT", bg_hex="F0FDF4", border_color="16A34A"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, bg_hex)
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        
        # Border
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        
        r_title = p.add_run(f"📌 {title}: ")
        r_title.bold = True
        r_title.font.size = Pt(10)
        r_title.font.color.rgb = RGBColor(22, 101, 52) if border_color == "16A34A" else PRIMARY_RGB
        
        r_text = p.add_run(text)
        r_text.font.size = Pt(9.5)
        r_text.font.color.rgb = DARK_TEXT_RGB
        
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # -------------------------------------------------------------
    # Title & Metadata
    # -------------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("Material Management, 22-Digit SKU Engine & Shopify Workflow Guide")
    run_title.bold = True
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = PRIMARY_RGB

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(14)
    run_sub = p_sub.add_run("End-to-End Operational Manual: Central Database, Automatic SKU Derivation, Inbound Shopify Webhook, Production Sheet Enrichment & Quick-Lookup")
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = SECONDARY_RGB

    # Meta box
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        [("System Version", "Phase 1 Complete (€1,750 Package)"), ("Target Audience", "Production Team, Admin Managers, Technical Team")],
        [("Live Application URL", "https://shopify-3d-viewersss-main.vercel.app"), ("Last Updated", "August 2026 / Version 2.0 Production Ready")]
    ]
    for row_idx, row in enumerate(meta_table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.width = Inches(3.25)
            set_cell_background(cell, BG_LIGHT_HEX)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            k, v = meta_data[row_idx][col_idx]
            r1 = p.add_run(f"{k}: ")
            r1.bold = True
            r1.font.size = Pt(9)
            r1.font.color.rgb = PRIMARY_RGB
            r2 = p.add_run(v)
            r2.font.size = Pt(9)
            r2.font.color.rgb = DARK_TEXT_RGB

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # -------------------------------------------------------------
    # 1. Executive Summary
    # -------------------------------------------------------------
    h1 = doc.add_heading("1. Executive Summary & Architecture Overview", level=1)
    h1.runs[0].font.color.rgb = PRIMARY_RGB
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph(
        "This document details the complete technical and operational workflow for the centralized Material Management System, "
        "the 22-digit SKU Generator, the Shopify Inbound Webhook connector, the enriched Production Sheet PDF/CSV export engine, "
        "and the SKU Quick-Lookup tool."
    )
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)

    p2 = doc.add_paragraph(
        "The core objective of this architecture is to establish a single source of truth for all technical material properties "
        "(fabric compositions, suppliers, exact yarn constructions, GSM weights, multiple finishings, and structured colour codes). "
        "When orders are placed on Shopify, the 22-digit SKU seamlessly bridges customer orders directly to factory cutting & production sheets "
        "with zero manual data re-entry."
    )
    p2.paragraph_format.line_spacing = 1.15
    p2.paragraph_format.space_after = Pt(8)

    add_callout(
        "The unified workflow: Create Material in Database (generates 6-digit ID) ➔ Select Material in SKU Generator (builds 22-digit SKU) "
        "➔ Enter SKU in Shopify Product Variant ➔ Customer buys on Shopify ➔ Webhook auto-creates Master & Sub-orders ➔ Production Sheet PDF/CSV auto-populates all fabric specs.",
        title="CORE AUTOMATION PRINCIPLE"
    )

    # -------------------------------------------------------------
    # 2. 22-Digit SKU Breakdown
    # -------------------------------------------------------------
    h2 = doc.add_heading("2. The 22-Digit SKU Structure", level=1)
    h2.runs[0].font.color.rgb = PRIMARY_RGB
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph(
        "The system employs a standardized 22-digit SKU format consisting of 8 defined segments separated by dashes (human format) "
        "or concatenated continuously (machine/barcode format):"
    )
    p.paragraph_format.space_after = Pt(6)

    # SKU Table
    sku_table = doc.add_table(rows=9, cols=5)
    sku_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Seg #", "Segment Key", "Segment Name", "Width", "Example Code & Meaning"]
    widths = [Inches(0.6), Inches(1.3), Inches(1.5), Inches(0.6), Inches(2.5)]

    # Header row
    for col_idx, text in enumerate(headers):
        cell = sku_table.cell(0, col_idx)
        cell.width = widths[col_idx]
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 3] else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)

    sku_rows = [
        ("1", "target_group", "Target Group", "1 digit", "1 = Men, 2 = Women, 3 = Unisex, 4 = Children"),
        ("2", "product_category", "Product Category", "2 digits", "01 = Shirt, 02 = Chino, 03 = Trousers, 04 = Belt"),
        ("3", "fabric_family", "Fabric Family", "2 digits", "01 = 100% Cotton, 02 = Cotton-Linen, 03 = Wool-Blend"),
        ("4", "fabric_type", "Fabric Type", "2 digits", "01 = Oxford, 02 = Poplin, 03 = Twill, 04 = Dobby"),
        ("5", "supplier", "Supplier Code", "3 digits", "005 = Albini, 021 = Monti, 034 = Safe Chino Mill"),
        ("6", "our_colour", "Our Colour Code", "3 digits", "143 = Night Blue (from 22 Colour Families 010–229)"),
        ("7", "reserved", "Reserved / Future", "3 digits", "000 = Standard Reserved Buffer"),
        ("8", "material_spec_id", "Material Spec ID", "6 digits", "000001 = Permanent Database ID pointing to Material Record")
    ]

    for row_idx, data in enumerate(sku_rows, start=1):
        for col_idx, val in enumerate(data):
            cell = sku_table.cell(row_idx, col_idx)
            cell.width = widths[col_idx]
            bg = BG_LIGHT_HEX if row_idx % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 3] else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.size = Pt(8.5)
            r.font.color.rgb = DARK_TEXT_RGB
            if col_idx == 0 or col_idx == 3:
                r.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Example display
    p_ex = doc.add_paragraph()
    p_ex.paragraph_format.space_before = Pt(4)
    p_ex.paragraph_format.space_after = Pt(8)
    r = p_ex.add_run("Example SKU (Human Format): ")
    r.bold = True
    r.font.color.rgb = PRIMARY_RGB
    r2 = p_ex.add_run("1-01-01-01-005-143-000-000123\n")
    r2.bold = True
    r2.font.color.rgb = SECONDARY_RGB
    r3 = p_ex.add_run("Example SKU (Machine / Barcode Format): ")
    r3.bold = True
    r3.font.color.rgb = PRIMARY_RGB
    r4 = p_ex.add_run("1010101005143000000123 (Exactly 22 continuous digits)")
    r4.bold = True

    # -------------------------------------------------------------
    # 3. Master Taxonomies: 20 Finishings & 22 Colour Families
    # -------------------------------------------------------------
    h3 = doc.add_heading("3. Master Taxonomies & Central Registry", level=1)
    h3.runs[0].font.color.rgb = PRIMARY_RGB
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph(
        "To ensure clean, un-duplicated data across suppliers, the database pre-seeds two comprehensive taxonomies:"
    )
    p.paragraph_format.space_after = Pt(4)

    # Finishing registry
    h3_sub1 = doc.add_heading("A. Multi-Select Finishing Master (20 Standard Finishings)", level=2)
    h3_sub1.runs[0].font.color.rgb = SECONDARY_RGB
    h3_sub1.paragraph_format.space_before = Pt(6)
    h3_sub1.paragraph_format.space_after = Pt(4)

    fin_p = doc.add_paragraph(
        "Materials can have multiple simultaneous finishes selected via checkboxes. The system stores these in the join table "
        "and automatically prints them as comma-separated labels on Production Sheets.\n"
        "1. Water repellent | 2. Easy care | 3. Easy ironing | 4. Non-iron | 5. Wrinkle resistant / Crease resistant | "
        "6. Anti-shrink / Shrink resistant | 7. Sanforized / Pre-shrunk | 8. Anti-pilling | 9. Anti-static | 10. Anti-bacterial / Antimicrobial | "
        "11. Anti-odour / Odour control | 12. Moisture wicking | 13. Quick dry | 14. UV protection / UPF | 15. Soft finish | "
        "16. Brushed | 17. Peached / Peach finish | 18. Enzyme finish | 19. Mercerized | 20. Washed finish"
    )
    fin_p.paragraph_format.line_spacing = 1.15
    fin_p.paragraph_format.space_after = Pt(6)

    # Colour taxonomy
    h3_sub2 = doc.add_heading("B. 3-Digit Colour Master (22 Colour Families, Ranges 010–229)", level=2)
    h3_sub2.runs[0].font.color.rgb = SECONDARY_RGB
    h3_sub2.paragraph_format.space_before = Pt(6)
    h3_sub2.paragraph_format.space_after = Pt(4)

    col_p = doc.add_paragraph(
        "Our Colour codes are strictly grouped into 22 families so shades can be organized systematically:\n"
        "• White Tones (010–019) | • Ivory / Cream / Ecru (020–029) | • Beige / Sand / Stone (030–039) | • Camel / Tan (040–049)\n"
        "• Brown Tones (050–059) | • Yellow Tones (060–069) | • Orange Tones (070–079) | • Red Tones (080–089)\n"
        "• Burgundy / Wine (090–099) | • Pink / Rose (100–109) | • Purple / Lilac (110–119) | • Light Blue Tones (120–129)\n"
        "• Medium / Bright Blue (130–139) | • Navy Tones (140–149) | • Turquoise / Teal / Petrol (150–159) | • Green Tones (160–169)\n"
        "• Olive / Khaki (170–179) | • Grey Tones (180–189) | • Charcoal / Anthracite (190–199) | • Black (200–209)\n"
        "• Metallic (210–219) | • Multicolour / Print (220–229)"
    )
    col_p.paragraph_format.line_spacing = 1.15
    col_p.paragraph_format.space_after = Pt(8)

    # -------------------------------------------------------------
    # 4. Step-by-Step Operating Guide
    # -------------------------------------------------------------
    h4 = doc.add_heading("4. Step-by-Step Operating Instructions", level=1)
    h4.runs[0].font.color.rgb = PRIMARY_RGB
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(4)

    steps = [
        ("Step 1: Adding a Material Specification",
         "URL: https://shopify-3d-viewersss-main.vercel.app/admin/materials",
         "1. Navigate to Admin ➔ Materials.\n"
         "2. Click '+ Add Material'.\n"
         "3. Select the Supplier from the dropdown (e.g., Albini, Monti).\n"
         "4. Enter the Supplier's Article / Quality Number (e.g., 'ALB-9842').\n"
         "5. Enter the Supplier's exact Colour Code & Name (e.g., 'C12 - Royal Sky').\n"
         "6. Select 'Our Colour' from the structured dropdown (e.g., '143 - Night Blue').\n"
         "7. Enter the Fabric Type, Composition (e.g., '100% Egyptian Cotton'), Width, Weight GSM (e.g., '120'), and Construction (e.g., 'Twill 2/1 80/2').\n"
         "8. Check all applicable Finishings checkboxes (e.g., Easy Care + Peached).\n"
         "9. Click 'Save Material'. The system automatically generates a unique 6-digit zero-padded Spec ID (e.g., '000001')."),

        ("Step 2: Generating the 22-Digit SKU in the Registry",
         "URL: https://shopify-3d-viewersss-main.vercel.app/admin/article-codes",
         "1. Navigate to Admin ➔ Article Codes.\n"
         "2. Click '+ New SKU'.\n"
         "3. Select Target Group (e.g., Men) and Product Category (e.g., Shirt).\n"
         "4. Select Fabric Family & Fabric Type.\n"
         "5. Select the Supplier ➔ The 'Material Specification' picker automatically filters to show materials available from that supplier.\n"
         "6. Select the Material Specification: The system instantly auto-fills the 6-digit Spec ID, Our Colour code, and Supplier code, locking them to prevent mismatch.\n"
         "7. The 'Reserved' field defaults to '000'.\n"
         "8. Review the live 22-digit preview and click 'Save SKU'. The SKU is now registered in the central database."),

        ("Step 3: Entering the SKU into Shopify Products",
         "Location: Shopify Admin ➔ Products ➔ Variants",
         "1. Open your Shopify Admin store.\n"
         "2. Open the specific Product or Variant corresponding to this garment/fabric.\n"
         "3. Paste the generated 22-digit SKU (e.g., '1-01-01-01-005-143-000-000123') into the variant's 'SKU' field.\n"
         "4. Save the product in Shopify.\n"
         "5. No further configuration is needed! The product is now linked to your central technical database."),

        ("Step 4: Customer Order ➔ Automatic Inbound Webhook Execution",
         "Automatic Background Workflow",
         "1. When a customer purchases the item on Shopify and completes payment, Shopify fires an 'Order payment' webhook to:\n"
         "   https://shopify-3d-viewersss-main.vercel.app/api/shopify/webhooks/orders-paid\n"
         "2. The webhook verifies the HMAC cryptographic signature using SHOPIFY_WEBHOOK_SECRET.\n"
         "3. Idempotency Check: Prevents duplicate orders if Shopify retries the webhook.\n"
         "4. Customer is auto-upserted with their shipping address and contact details.\n"
         "5. Master Order is created with origin = 'shopify'.\n"
         "6. Sub-orders are spawned for each garment item in the order.\n"
         "7. The 22-digit SKU is parsed to extract the 6-digit Material Spec ID, stamping the human & machine article codes onto the sub-order."),

        ("Step 5: Production Sheet PDF & Factory CSV Export",
         "URL: https://shopify-3d-viewersss-main.vercel.app/admin/orders",
         "1. In Admin ➔ Orders, open any order and click 'Export Production Sheet (PDF)' or 'Export CSV'.\n"
         "2. The PDF generator reads the embedded Material Spec ID from the SKU and fetches the live database record.\n"
         "3. Below the Article Code Barcode, a dedicated 'Material Specification' section is printed with:\n"
         "   • Supplier & Supplier Article Number\n"
         "   • Supplier Colour Code & Name\n"
         "   • Fabric Composition & Weave Construction\n"
         "   • Fabric Weight (GSM) & Width\n"
         "   • Active Technical Finishings list\n"
         "4. The CSV export includes all discrete material columns, ready for production batching."),

        ("Step 6: Real-Time SKU Quick-Lookup Tool",
         "URL: https://shopify-3d-viewersss-main.vercel.app/admin/sku-lookup",
         "1. Navigate to Admin ➔ SKU Lookup.\n"
         "2. Paste or barcode-scan ANY 22-digit SKU (either '1-01-01-01-005-143-000-000123' or '1010101005143000000123').\n"
         "3. The tool instantly decodes all 8 segments into human-readable labels.\n"
         "4. An interactive Material Specification Card renders below, displaying the full linked technical record, finishings tags, and direct export actions.")
    ]

    for title, subtitle, content in steps:
        h = doc.add_heading(title, level=2)
        h.runs[0].font.color.rgb = PRIMARY_RGB
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(2)
        
        p_sub = doc.add_paragraph(subtitle)
        p_sub.runs[0].font.size = Pt(9.5)
        p_sub.runs[0].font.bold = True
        p_sub.runs[0].font.color.rgb = SECONDARY_RGB
        p_sub.paragraph_format.space_after = Pt(4)
        
        p_body = doc.add_paragraph(content)
        p_body.paragraph_format.line_spacing = 1.15
        p_body.paragraph_format.space_after = Pt(8)

    # -------------------------------------------------------------
    # 5. Database Schema Reference
    # -------------------------------------------------------------
    h5 = doc.add_heading("5. Database Schema & Tables Reference", level=1)
    h5.runs[0].font.color.rgb = PRIMARY_RGB
    h5.paragraph_format.space_before = Pt(12)
    h5.paragraph_format.space_after = Pt(4)

    schema_table = doc.add_table(rows=6, cols=3)
    schema_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    s_headers = ["Table Name", "Primary Key & Indexes", "Purpose & Key Columns"]
    s_widths = [Inches(2.0), Inches(1.8), Inches(2.7)]

    for col_idx, text in enumerate(s_headers):
        cell = schema_table.cell(0, col_idx)
        cell.width = s_widths[col_idx]
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)

    s_rows = [
        ("material_specifications", "id (UUID PK)\nspec_id (UNIQUE)\nidx_matspec_supplier\nidx_matspec_colour", "Central technical record. spec_id auto-generated via sequence ('000001'). Stores supplier_article_number, supplier_colour_name, composition, weight_gsm, construction."),
        ("finishing_master", "id (UUID PK)\ncode (UNIQUE 2-digit)", "Registry of 20 standard finishings ('01' to '20'). Seeded with permanent labels and sort order."),
        ("material_specification_finishings", "(material_spec_id, finishing_id) Composite PK", "Join table for multi-select finishings. Cascades on delete to maintain data integrity."),
        ("colour_master", "id (UUID PK)\ncode (UNIQUE 3-digit)", "Structured taxonomy of 22 colour families (010–229) with family ranges, sub-colour labels, and sort keys."),
        ("product_skus", "id (UUID PK)\nsku_key (UNIQUE)\nidx_skus_category", "Registry of 22-digit SKUs with 8 discrete segment codes, our_colour_code, material_spec_id, article_human, and article_machine.")
    ]

    for row_idx, (tname, idxs, desc) in enumerate(s_rows, start=1):
        r = schema_table.rows[row_idx]
        bg = BG_LIGHT_HEX if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate([tname, idxs, desc]):
            cell = r.cells[col_idx]
            cell.width = s_widths[col_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(8.5)
            run.font.color.rgb = DARK_TEXT_RGB
            if col_idx == 0:
                run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # -------------------------------------------------------------
    # 6. Technical URLs & Endpoints Summary
    # -------------------------------------------------------------
    h6 = doc.add_heading("6. Complete URLs & API Endpoints Directory", level=1)
    h6.runs[0].font.color.rgb = PRIMARY_RGB
    h6.paragraph_format.space_before = Pt(12)
    h6.paragraph_format.space_after = Pt(4)

    url_table = doc.add_table(rows=7, cols=3)
    url_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    u_headers = ["Endpoint / Page", "Method & Auth", "Description"]
    u_widths = [Inches(2.5), Inches(1.3), Inches(2.7)]

    for col_idx, text in enumerate(u_headers):
        cell = url_table.cell(0, col_idx)
        cell.width = u_widths[col_idx]
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)

    u_rows = [
        ("/admin/materials", "UI (Admin Auth)", "Central Material Specification Management table with search, filters, multi-select finishings, and CSV export."),
        ("/admin/article-codes", "UI (Admin Auth)", "22-digit SKU Registry with smart dropdown builder, auto-fill from material specs, and live preview."),
        ("/admin/sku-lookup", "UI (Admin Auth)", "Real-time SKU decoder. Paste/scan any 22-digit code to view segment meanings and material record."),
        ("/api/shopify/webhooks/orders-paid", "POST (HMAC Secret)", "Shopify Inbound Webhook. Receives paid orders, creates customer/orders/sub-orders, parses SKUs."),
        ("/api/admin/materials/export", "GET (Admin)", "Direct download of complete Material Specifications database as CSV."),
        ("/api/admin/article-codes/export", "GET (Admin)", "Direct download of complete 22-digit SKU Registry with decoded names as CSV.")
    ]

    for row_idx, (ep, method, desc) in enumerate(u_rows, start=1):
        r = url_table.rows[row_idx]
        bg = BG_LIGHT_HEX if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate([ep, method, desc]):
            cell = r.cells[col_idx]
            cell.width = u_widths[col_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(8.5)
            run.font.color.rgb = DARK_TEXT_RGB
            if col_idx == 0:
                run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # 7. Quality Assurance & Sign-Off
    # -------------------------------------------------------------
    h7 = doc.add_heading("7. Production Sign-Off & Status", level=1)
    h7.runs[0].font.color.rgb = PRIMARY_RGB
    h7.paragraph_format.space_before = Pt(12)
    h7.paragraph_format.space_after = Pt(4)

    add_callout(
        "All components for the €1,750 Material Management package have passed automated TypeScript compilation (pnpm build), "
        "database migration execution, HMAC webhook secret configuration, and live production deployment at "
        "https://shopify-3d-viewersss-main.vercel.app. The system is 100% operational and ready for live production orders.",
        title="VERIFICATION & DEPLOYMENT COMPLETE"
    )

    output_path = "/Users/hasan/Documents/Aenfinite Projecrts/shopify-3d-viewersss-main/shopify-3d-viewersss-main/MATERIAL-MANAGEMENT-AND-SHOPIFY-WORKFLOW-GUIDE.docx"
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    create_document()
